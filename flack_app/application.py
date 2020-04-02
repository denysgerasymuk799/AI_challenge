import json
import os

from docx import Document
from flask import Flask, render_template, url_for, request, redirect, session, send_from_directory, send_file

app = Flask(__name__)
app.secret_key = "ylkv0bCqPliokdenmvtcTtx19gVnGBsL"


def get_courses(current_skills):
    """
    test function
    :param skills:
    :return:
    """
    # with open(os.path.join(os.getcwd(), "courses_for_profession.json"),
    #           encoding='utf-8') as f:
    #     course_list = json.load(f)

    with open(os.path.join(os.getcwd(), 'user_data', 'user_data.json'), 'r', encoding='utf-8') as \
            user_data_json_from_file:
        all_user_data = json.load(user_data_json_from_file)

    title_profession = '+'.join(all_user_data['profession'].split())
    courses_list_input_profession = {}
    with open(os.path.join(os.getcwd(), 'user_data', title_profession + '.json'), 'r', encoding='utf-8') as \
            json_file:
        course_list = json.load(json_file)
        for skill in course_list.keys():
            if skill not in current_skills:
                courses_list_input_profession[skill] = course_list[skill]

    id = 0
    for i, skill in enumerate(courses_list_input_profession.keys()):
        for j, course in enumerate(courses_list_input_profession[skill].keys()):
            courses_list_input_profession[skill][course]["name"] = course
            courses_list_input_profession[skill][course]["id"] = "course-" + str(id)
            id += 1
    return courses_list_input_profession


def skills_for_job(job):
    """
    Test function
    :param job:
    :return:
    """

    with open(os.path.join(os.getcwd(), 'static', 'filtered_skills_for_professions.json'), 'r', encoding='utf-8') as \
            json_file:
        all_skills = json.load(json_file)

    title_profession = '+'.join(job.split())
    return all_skills[title_profession]


global skills
global courses


@app.route("/a")
def a():
    return render_template("skills.html")


@app.route('/', methods=['POST', 'GET'])
def start():
    if request.method == 'POST':
        job = request.values.get("job").lower()
        eng_job_titles = ["system administrator", "analyst"]
        ru_job_titles = ["системный администратор", "аналитик"]
        if job in eng_job_titles:
            job = ru_job_titles[eng_job_titles.index(job)]

        global skills

        with open(os.path.join(os.getcwd(), 'user_data', 'user_data.json'), 'w', encoding='utf-8') as \
                user_data_json_from_file:
            user_data_json = {}
            user_data_json['profession'] = job
            json.dump(user_data_json, user_data_json_from_file, indent=4, ensure_ascii=False)

        skills = skills_for_job(job)  # your function
        return redirect(url_for("middle"))
    else:
        return render_template("request.html")


@app.route('/skills', methods=['POST', 'GET'])
def middle():
    if request.method == 'POST':
        current_skills = request.form.getlist("chosen_skills")
        print(current_skills)
        global courses
        courses = get_courses(current_skills)  # your function

        return redirect(url_for("index"))
    else:
        return render_template("skills.html", skills=skills)


@app.route('/courses', methods=['POST', 'GET'])
def index():
    return render_template("one_section.html", courses_list=courses)


@app.route('/selected', methods=['POST', 'GET'])
def selected():
    my_courses = []
    data = request.form
    session['my_courses'] = list(data.keys())
    for i, skill in enumerate(courses.keys()):
        for j, course in enumerate(courses[skill].keys()):
            if courses[skill][course]['id'] in data:
                my_courses.append(courses[skill][course])
    print(my_courses)
    return render_template("selected.html", my_courses=my_courses)


@app.route("/download")
def download():
    data = session['my_courses']
    document = Document()

    for i, skill in enumerate(courses.keys()):
        for j, course in enumerate(courses[skill].keys()):
            if courses[skill][course]['id'] in data:
                document.add_heading(courses[skill][course]['name'], level=2)
                document.add_paragraph("URL: " + courses[skill][course]['url'])
                document.add_paragraph(courses[skill][course]['long_description'])
    path = 'data/selected.docx'
    document.save(path)
    return send_file(path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
